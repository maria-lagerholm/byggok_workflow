import pandas as pd
import os
import re
from docx import Document
from docx.shared import Pt

# Get the current working directory
base_dir = os.getcwd()

# Path to the Excel file
file_path = os.path.join(base_dir, 'kunder', 'kundregister.xlsx')

# Read property designations and create modified list
prop_desig = pd.read_excel(file_path)['Fastighetsbeteckning'].tolist()
prop_desig_dir = [p.replace(':', '_').replace(' ', '_') for p in prop_desig]

# Function to find original property designation
def find_original_prop_desig(dir_name):
    return prop_desig[prop_desig_dir.index(dir_name)]

# Function to set font to Arial 11
def set_arial_11(run):
    run.font.name = 'Arial'
    run.font.size = Pt(11)

# Function to extract all text in a paragraph (including across multiple runs)
def get_full_text(paragraph):
    return ''.join(run.text for run in paragraph.runs)

# Function to clear all runs in a paragraph
def clear_paragraph(paragraph):
    for run in paragraph.runs:
        p = run._element
        p.getparent().remove(p)

# Function to update text in a paragraph
def update_text(paragraph, original_prop_desig):
    pattern = r'Fastighetsbeteckning:\s*([^\n]*)'
    
    # Get full text in the paragraph (combining all runs)
    full_text = get_full_text(paragraph)
    
    # Check for matches in the full text (including handling empty fields or extra spaces)
    match = re.search(pattern, full_text, re.IGNORECASE)
    
    if match:
        # Get the text that follows 'Fastighetsbeteckning:'
        current_value = match.group(1).strip()
        
        # Only update if the current value is empty or different
        if not current_value or current_value != original_prop_desig:
            # Clear the paragraph
            clear_paragraph(paragraph)
            
            # Add the updated text
            new_text = f"Fastighetsbeteckning: {original_prop_desig}"
            run = paragraph.add_run(new_text)
            set_arial_11(run)
            return True
    return False

# Function to process headers and footers
def process_header_footer(part, original_prop_desig):
    modified = False
    for paragraph in part.paragraphs:
        if update_text(paragraph, original_prop_desig):
            modified = True
    for table in part.tables:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if "Fastighetsbeteckning:" in cell.text and j + 1 < len(row.cells):
                    next_cell = row.cells[j + 1]
                    if next_cell.text.strip() != original_prop_desig:
                        next_cell.text = original_prop_desig
                        set_arial_11(next_cell.paragraphs[0].runs[0])
                        modified = True
    return modified

# Iterate through each directory
for current_dir in prop_desig_dir:
    dir_path = os.path.join(base_dir, 'kunder', current_dir)
    if not os.path.exists(dir_path):
        print(f"Directory '{dir_path}' does not exist. Skipping.")
        continue

    original_prop_desig = find_original_prop_desig(current_dir)
    print(f"\nProcessing directory: {current_dir}")

    for file in os.listdir(dir_path):
        if file.endswith('.docx'):
            doc_path = os.path.join(dir_path, file)
            doc = Document(doc_path)
            modified = False

            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for j, cell in enumerate(row.cells):
                        if "Fastighetsbeteckning:" in cell.text and j + 1 < len(row.cells):
                            next_cell = row.cells[j + 1]
                            if next_cell.text.strip() != original_prop_desig:
                                next_cell.text = original_prop_desig
                                set_arial_11(next_cell.paragraphs[0].runs[0])
                                modified = True

            # Check paragraphs (including captions)
            for paragraph in doc.paragraphs:
                if update_text(paragraph, original_prop_desig):
                    modified = True

            # Check headers and footers
            for section in doc.sections:
                modified |= process_header_footer(section.header, original_prop_desig)
                modified |= process_header_footer(section.first_page_header, original_prop_desig)
                modified |= process_header_footer(section.footer, original_prop_desig)
                modified |= process_header_footer(section.first_page_footer, original_prop_desig)

            if modified:
                doc.save(doc_path)
                print(f"Updated {file}")
            else:
                print(f"No changes needed for {file}")

print("\nProcessing of all directories is complete.")