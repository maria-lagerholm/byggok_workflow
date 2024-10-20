import pandas as pd
import os
import sys
import re
from docx import Document
from docx.shared import Pt

def find_original_prop_desig(dir_name, prop_desig, prop_desig_dir):
    return prop_desig[prop_desig_dir.index(dir_name)]

def set_arial_11(run):
    run.font.name = 'Arial'
    run.font.size = Pt(11)

def get_full_text(paragraph):
    return ''.join(run.text for run in paragraph.runs)

def clear_paragraph(paragraph):
    for run in paragraph.runs:
        p = run._element
        p.getparent().remove(p)

def update_text(paragraph, original_prop_desig):
    pattern = r'Fastighetsbeteckning:\s*([^\n]*)'
    full_text = get_full_text(paragraph)
    match = re.search(pattern, full_text, re.IGNORECASE)
    
    if match:
        current_value = match.group(1).strip()
        if not current_value or current_value != original_prop_desig:
            clear_paragraph(paragraph)
            new_text = f"Fastighetsbeteckning: {original_prop_desig}"
            run = paragraph.add_run(new_text)
            set_arial_11(run)
            return True
    return False

def process_header_footer(part, original_prop_desig):
    modified = False
    for paragraph in part.paragraphs:
        if update_text(paragraph, original_prop_desig):
            modified = True
    for table in part.tables:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if "Fastighetsbeteckning:" in cell.text and j + 1 < len(row.cells):
                    next_cell = row.cells[j + 1]
                    if next_cell.text.strip() != original_prop_desig:
                        next_cell.text = original_prop_desig
                        set_arial_11(next_cell.paragraphs[0].runs[0])
                        modified = True
    return modified

def main():
    if getattr(sys, 'frozen', False):
        base_dir = sys._MEIPASS
    else:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    file_path = os.path.join(base_dir, 'kunder', 'kundregister.xlsx')
    print(f"Reading Excel file from: {file_path}")

    prop_desig = pd.read_excel(file_path)['Fastighetsbeteckning'].tolist()
    prop_desig_dir = [p.replace(':', '_').replace(' ', '_') for p in prop_desig]

    total_dirs = len(prop_desig_dir)
    processed_dirs = 0
    updated_files = 0

    for current_dir in prop_desig_dir:
        processed_dirs += 1
        dir_path = os.path.join(base_dir, 'kunder', current_dir)
        if not os.path.exists(dir_path):
            continue

        original_prop_desig = find_original_prop_desig(current_dir, prop_desig, prop_desig_dir)

        for file in os.listdir(dir_path):
            if file.endswith('.docx'):
                doc_path = os.path.join(dir_path, file)
                doc = Document(doc_path)
                modified = False

                for table in doc.tables:
                    for row in table.rows:
                        for j, cell in enumerate(row.cells):
                            if "Fastighetsbeteckning:" in cell.text and j + 1 < len(row.cells):
                                next_cell = row.cells[j + 1]
                                if next_cell.text.strip() != original_prop_desig:
                                    next_cell.text = original_prop_desig
                                    set_arial_11(next_cell.paragraphs[0].runs[0])
                                    modified = True

                for paragraph in doc.paragraphs:
                    if update_text(paragraph, original_prop_desig):
                        modified = True

                for section in doc.sections:
                    modified |= process_header_footer(section.header, original_prop_desig)
                    modified |= process_header_footer(section.first_page_header, original_prop_desig)
                    modified |= process_header_footer(section.footer, original_prop_desig)
                    modified |= process_header_footer(section.first_page_footer, original_prop_desig)

                if modified:
                    doc.save(doc_path)
                    print(f"Updated: {file} in {current_dir}")
                    updated_files += 1

        # Print progress every 10% of directories processed
        if processed_dirs % max(1, total_dirs // 10) == 0:
            print(f"Progress: {processed_dirs}/{total_dirs} directories processed")

    print(f"\nPart 2 completed.")
    print(f"Updated {updated_files} files across {total_dirs} directories.")
    print("\nStarting next script (Part 3)...")

if __name__ == "__main__":
    main()