import pandas as pd
import os
import sys
import re
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
import urllib.parse

def find_original_info(dir_name, prop_desig_dir, info_lists, columns):
    index = prop_desig_dir.index(dir_name)
    info = {col: info_lists[col][index] for col in columns}
    if pd.notna(info['Besiktningsdag']):
        info['Besiktningsdag'] = info['Besiktningsdag'].strftime('%Y-%m-%d')
    if pd.notna(info['Klockan']):
        info['Klockan'] = info['Klockan'].strftime('%H:%M')
    return {k: str(v).strip() if pd.notna(v) else '' for k, v in info.items()}

def set_arial_11(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

def should_update(current_value, original_value):
    return not current_value.strip() or current_value.strip().lower() == 'ange adress'

def process_tables(tables, original_info):
    modified = False
    for table in tables:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                for col in original_info.keys():
                    if re.search(f"{re.escape(col)}:(?!_)", cell.text.strip(), re.IGNORECASE) and j + 1 < len(row.cells):
                        next_cell = row.cells[j + 1]
                        original_value = original_info[col]
                        current_value = next_cell.text.strip()
                        if should_update(current_value, original_value):
                            next_cell.text = original_value
                            set_arial_11(next_cell)
                            modified = True
    return modified

def process_header_footer(part, original_info):
    modified = False
    for paragraph in part.paragraphs:
        for col in original_info.keys():
            match = re.search(f"{re.escape(col)}:(?!_)\s*(.*)", paragraph.text.strip(), re.IGNORECASE)
            if match:
                current_value = match.group(1).strip()
                original_value = original_info[col]
                if should_update(current_value, original_value):
                    new_text = re.sub(f"{re.escape(col)}:(?!_)\s*(.*)", f"{col}: {original_value}", paragraph.text, flags=re.IGNORECASE)
                    paragraph.text = new_text
                    set_arial_11(paragraph)
                    modified = True
    for table in part.tables:
        modified |= process_tables([table], original_info)
    return modified

def create_google_calendar_url(info):
    event_name = f"Besiktning {info['Kommun']}"
    start_time = datetime.strptime(f"{info['Besiktningsdag']} {info['Klockan']}", "%Y-%m-%d %H:%M")
    end_time = start_time + timedelta(hours=2)
    
    start_str = start_time.strftime("%Y%m%dT%H%M%S")
    end_str = end_time.strftime("%Y%m%dT%H%M%S")
    
    details = f"{info['Fastighetsägare']}\n"
    details += f"{info['Telefon']}\n"
    details += f"{info['E-post']}\n"
    
    base_url = "https://www.google.com/calendar/render?action=TEMPLATE"
    event_params = {
        "text": event_name,
        "dates": f"{start_str}/{end_str}",
        "details": details,
        "location": info['Adress'],
        "reminders": "POPUP,1440",
    }
    
    url = f"{base_url}&{urllib.parse.urlencode(event_params)}"
    return url

def main():
    if getattr(sys, 'frozen', False):
        base_dir = sys._MEIPASS
    else:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    file_path = os.path.join(base_dir, 'kunder', 'kundregister.xlsx')
    print(f"Reading Excel file from: {file_path}")

    df = pd.read_excel(file_path)
    df['Besiktningsdag'] = pd.to_datetime(df['Besiktningsdag'], format='%Y-%m-%d', errors='coerce')
    df['Klockan'] = pd.to_datetime(df['Klockan'], format='%H:%M:%S', errors='coerce')

    columns = ['Adress', 'Kommun', 'Fastighetsägare', 'Uppdragsgivare', 'Postadress', 'E-post', 'Telefon', 
               'Uppdragsnummer', 'Besiktningsdag', 'Klockan', 'Kostnad']
    info_lists = {col: df[col].tolist() for col in columns}
    prop_desig = df['Fastighetsbeteckning'].tolist()
    prop_desig_dir = [p.replace(':', '_').replace(' ', '_') for p in prop_desig]

    total_dirs = len(prop_desig_dir)
    processed_dirs = 0
    updated_files = 0
    created_calendar_urls = 0

    for current_dir in prop_desig_dir:
        processed_dirs += 1
        dir_path = os.path.join(base_dir, 'kunder', current_dir)
        if not os.path.exists(dir_path):
            continue

        original_info = find_original_info(current_dir, prop_desig_dir, info_lists, columns)
        
        url_file_path = os.path.join(dir_path, f"{current_dir}_besiktning_google_calendar.url")
        if not os.path.exists(url_file_path):
            calendar_url = create_google_calendar_url(original_info)
            with open(url_file_path, 'w') as f:
                f.write("[InternetShortcut]\n")
                f.write(f"URL={calendar_url}\n")
            print(f"Created new Google Calendar URL file: {url_file_path}")
            created_calendar_urls += 1

        for file in os.listdir(dir_path):
            if file.endswith('.docx'):
                doc_path = os.path.join(dir_path, file)
                doc = Document(doc_path)
                modified = False

                for table in doc.tables:
                    modified |= process_tables([table], original_info)

                for section in doc.sections:
                    modified |= process_header_footer(section.header, original_info)
                    modified |= process_header_footer(section.first_page_header, original_info)
                    modified |= process_header_footer(section.footer, original_info)
                    modified |= process_header_footer(section.first_page_footer, original_info)

                if modified:
                    doc.save(doc_path)
                    print(f"Updated: {file} in {current_dir}")
                    updated_files += 1

        # Print progress every 10% of directories processed
        if processed_dirs % max(1, total_dirs // 10) == 0:
            print(f"Progress: {processed_dirs}/{total_dirs} directories processed")

    print(f"\nProcessing complete. Updated {updated_files} files and created {created_calendar_urls} calendar URLs across {total_dirs} directories.")

if __name__ == "__main__":
    main()